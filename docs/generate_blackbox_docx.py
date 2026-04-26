"""
Generate BAB-BLACKBOX.docx dari PENGUJIAN-BLACKBOX.md
Standar: Pedoman Penulisan Skripsi FTI Unisba Blitar 2025-2026

Aturan FTI:
- Judul TABEL di ATAS tabel (§5.2.6) — berbeda dari gambar
- Font tabel: TNR 10pt, spasi 1
- Font isi: TNR 12pt, spasi 1.5
- Margin: Kiri 4cm, Atas/Kanan/Bawah 3cm
"""
import os
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(SCRIPT_DIR, "PENGUJIAN-BLACKBOX.md")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "BAB-BLACKBOX.docx")


def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)


def add_formatted_text(paragraph, text, size=12, bold=False, color=None):
    """Add text with italic support (*text*) to a paragraph."""
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size, italic=True, bold=bold, color=color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, bold=bold, color=color)


def add_cell_content(cell, text, size=10, bold=False, color=None):
    """Fill a table cell with formatted text, handling <br> as line breaks."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    lines = re.split(r'<br\s*/?>', text.strip())
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run('\n')
            set_font(run, size=size)
        line = line.strip().lstrip('- ').strip()
        add_formatted_text(p, line, size=size, bold=bold, color=color)


def setup_document():
    """Create document with FTI margins and default font."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    return doc


def add_heading(doc, text, level=2):
    """Add heading with TNR font."""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)
    return h


def add_para(doc, text, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add paragraph with italic support."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(1.27)
    add_formatted_text(p, text)
    return p


def add_precondition(doc, text):
    """Add precondition text (italic, no indent)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, italic=True)


def add_table_caption(doc, caption):
    """Add table title ABOVE the table (§5.2.6 FTI)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(caption)
    set_font(run, size=10, bold=True)


def create_test_table(doc, rows_data):
    """Create a 6-column test table with header and data rows."""
    headers = ['No', 'Skenario Pengujian', 'Langkah-langkah',
               'Hasil Diharapkan', 'Hasil Diperoleh', 'Hasil Pengujian']

    table = doc.add_table(rows=1 + len(rows_data), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths (total available: 14cm = 21 - 4 - 3)
    widths = [Cm(0.8), Cm(2.8), Cm(3.2), Cm(3.0), Cm(3.0), Cm(1.2)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(header_text)
        set_font(run, size=10, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if col_idx == 0:  # No column - center
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(str(cell_text))
                set_font(run, size=10)
            elif col_idx == 5:  # Hasil Pengujian - bold green
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(str(cell_text))
                set_font(run, size=10, bold=True, color=(0, 128, 0))
            else:
                add_cell_content(cell, str(cell_text), size=10)

    doc.add_paragraph()  # spacing after table
    return table


def parse_markdown_tables(filepath):
    """Parse PENGUJIAN-BLACKBOX.md and extract structured test tables only."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    tables = []  # List of {caption, precondition, rows}
    current_table = None
    current_rows = []
    header_seen = False
    in_test_section = False  # Only parse after "## 2. Tabel Pengujian"

    for line in content.split('\n'):
        line = line.rstrip('\r')

        # Only start parsing after "## 2. Tabel Pengujian"
        if line.startswith('## 2. Tabel Pengujian'):
            in_test_section = True
            continue

        # Stop parsing at "## 3. Rekapitulasi"
        if line.startswith('## 3. Rekapitulasi'):
            if current_rows and current_table:
                current_table['rows'] = current_rows
            break

        if not in_test_section:
            continue

        # Detect main table headers (### Tabel 4.X)
        if line.startswith('### Tabel '):
            if current_rows and current_table:
                current_table['rows'] = current_rows
            caption = line.replace('### ', '').strip()
            current_table = {'caption': caption, 'precondition': '', 'rows': []}
            tables.append(current_table)
            current_rows = []
            header_seen = False
            continue

        # Detect sub-section headers (#### a. Sub-Modul: ...)
        if line.startswith('#### '):
            if current_rows and current_table:
                current_table['rows'] = current_rows
            sub_caption = line.replace('#### ', '').strip()
            current_table = {'caption': sub_caption, 'precondition': '', 'rows': []}
            tables.append(current_table)
            current_rows = []
            header_seen = False
            continue

        # Detect precondition
        if line.startswith('> *Precondition') and current_table:
            pre_text = line.lstrip('> ').strip()
            pre_text = re.sub(r'^\*|\*$', '', pre_text).strip()
            current_table['precondition'] = pre_text
            continue

        # Detect table rows (6-column test tables only)
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(c.replace('-', '').strip() == '' for c in cells):
                header_seen = True
                continue
            if not header_seen:
                continue
            if len(cells) == 6:
                current_rows.append(cells)
            continue
        else:
            if current_rows and current_table:
                current_table['rows'] = current_rows
                current_rows = []
            header_seen = False

    return tables


def build_document(doc):
    """Build the complete blackbox testing document."""

    # ============================================================
    # 1. PENDAHULUAN
    # ============================================================
    add_heading(doc, '1. Pendahuluan', level=2)

    add_heading(doc, '1.1 Tujuan Pengujian', level=3)
    add_para(doc,
        'Dokumen ini bertujuan untuk mendokumentasikan hasil *Pengujian Black-Box* '
        '(*Black-Box Testing*) terhadap sistem POS PWA yang dikembangkan sebagai luaran '
        'penelitian tugas akhir. Pengujian dilakukan untuk memvalidasi bahwa seluruh '
        'fungsionalitas sistem berperilaku sesuai dengan spesifikasi kebutuhan yang telah '
        'ditetapkan, tanpa mempertimbangkan struktur kode internal.')

    add_heading(doc, '1.2 Ruang Lingkup', level=3)
    add_para(doc,
        'Pengujian mencakup 7 modul fungsional utama sistem POS PWA:')

    # Scope table
    add_table_caption(doc, 'Tabel 4.1 \u2014 Ruang Lingkup Modul Pengujian')
    scope_headers = ['Kode', 'Modul', 'Deskripsi']
    scope_data = [
        ['F-01', 'Autentikasi', 'Proses login, logout, proteksi rute, dan persistensi sesi'],
        ['F-02', 'Beranda (Home Hub)', 'Dasbor akses cepat dengan statistik ringkas dan status koneksi'],
        ['F-03', 'Dashboard Analitik', 'Grafik dan statistik penjualan dengan filter periode'],
        ['F-04', 'Kasir (Point of Sale)', 'Transaksi, keranjang, pembayaran, struk, dan dukungan offline'],
        ['F-05', 'Manajemen Produk', 'CRUD produk, kelola kategori, dan toggle status'],
        ['F-06', 'Laporan', 'Riwayat, filter transaksi, dan sinkronisasi offline'],
        ['F-07', 'Pengaturan', 'Info toko, printer Bluetooth, dan profil akun'],
    ]
    tbl_scope = doc.add_table(rows=1 + len(scope_data), cols=3)
    tbl_scope.style = 'Table Grid'
    tbl_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(scope_headers):
        cell = tbl_scope.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_font(run, size=10, bold=True)
        set_cell_shading(cell, 'D9E2F3')
    for r_idx, row in enumerate(scope_data):
        for c_idx, val in enumerate(row):
            cell = tbl_scope.rows[r_idx + 1].cells[c_idx]
            add_cell_content(cell, val, size=10)
    doc.add_paragraph()

    add_heading(doc, '1.3 Metodologi', level=3)
    add_para(doc,
        'Pengujian menggunakan metode *Black-Box Testing* dengan pendekatan:')
    # EP
    p_ep = doc.add_paragraph()
    p_ep.paragraph_format.line_spacing = 1.5
    p_ep.paragraph_format.left_indent = Cm(1.27)
    add_formatted_text(p_ep, '1. *Equivalence Partitioning* (EP): Membagi domain input ke dalam '
        'kelas-kelas ekuivalen untuk mengurangi jumlah kasus uji tanpa mengorbankan cakupan.')
    # BVA
    p_bva = doc.add_paragraph()
    p_bva.paragraph_format.line_spacing = 1.5
    p_bva.paragraph_format.left_indent = Cm(1.27)
    add_formatted_text(p_bva, '2. *Boundary Value Analysis* (BVA): Menguji nilai-nilai batas kritis '
        'pada setiap partisi input untuk mendeteksi kesalahan pada kondisi tepi.')

    add_para(doc, 'Kriteria kelulusan yang digunakan adalah sebagai berikut:', indent=True)
    p_pass = doc.add_paragraph()
    p_pass.paragraph_format.line_spacing = 1.5
    p_pass.paragraph_format.left_indent = Cm(1.27)
    run_b = p_pass.add_run('Berhasil')
    set_font(run_b, bold=True)
    run_t = p_pass.add_run(' \u2014 Sistem menghasilkan output yang identik dengan Hasil Diharapkan.')
    set_font(run_t)

    p_fail = doc.add_paragraph()
    p_fail.paragraph_format.line_spacing = 1.5
    p_fail.paragraph_format.left_indent = Cm(1.27)
    run_b2 = p_fail.add_run('Tidak Berhasil')
    set_font(run_b2, bold=True)
    run_t2 = p_fail.add_run(' \u2014 Sistem menghasilkan output yang berbeda dari Hasil Diharapkan.')
    set_font(run_t2)

    add_heading(doc, '1.4 Lingkungan Pengujian', level=3)
    add_table_caption(doc, 'Tabel 4.2 \u2014 Lingkungan Pengujian')
    env_data = [
        ['Browser', 'Google Chrome (versi terbaru)'],
        ['Mode', 'Desktop (\u22651024px) dan Mobile (375px)'],
        ['Koneksi', 'Online (WiFi) dan Offline (DevTools \u2192 Offline)'],
        ['Sistem Operasi', 'Windows 11'],
        ['URL Aplikasi', 'http://localhost:3000 (Development)'],
        ['Versi Aplikasi', 'v0.5.0'],
        ['Tanggal Pengujian', '2026-04-15'],
    ]
    tbl_env = doc.add_table(rows=len(env_data), cols=2)
    tbl_env.style = 'Table Grid'
    tbl_env.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(env_data):
        for c_idx, val in enumerate(row):
            cell = tbl_env.rows[r_idx].cells[c_idx]
            add_cell_content(cell, val, size=10, bold=(c_idx == 0))
    doc.add_paragraph()

    # ============================================================
    # 2. TABEL PENGUJIAN (parsed from markdown)
    # ============================================================
    add_heading(doc, '2. Tabel Pengujian', level=2)

    print("\nParsing PENGUJIAN-BLACKBOX.md...")
    tables = parse_markdown_tables(MD_FILE)

    table_number = 3  # Start from 4.3 (4.1=scope, 4.2=env)
    total_tc = 0

    for tbl_data in tables:
        caption = tbl_data.get('caption', '')
        rows = tbl_data.get('rows', [])
        precondition = tbl_data.get('precondition', '')

        if not rows:
            # Sub-section header without data rows
            if caption:
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(12)
                run_sub = p_sub.add_run(caption)
                set_font(run_sub, size=12, bold=True)
            continue

        # Replace Tabel numbering with sequential 4.X
        if caption.startswith('Tabel '):
            display_caption = re.sub(r'Tabel \d+\.\d+', f'Tabel 4.{table_number}', caption)
        else:
            display_caption = f'Tabel 4.{table_number} \u2014 {caption}'

        add_table_caption(doc, display_caption)

        if precondition:
            add_precondition(doc, precondition)

        create_test_table(doc, rows)

        tc_count = len(rows)
        total_tc += tc_count
        table_number += 1
        print(f"  [OK] {display_caption} ({tc_count} TC)")

    # ============================================================
    # 3. REKAPITULASI
    # ============================================================
    add_heading(doc, '3. Rekapitulasi Hasil Pengujian', level=2)

    recap_caption = f'Tabel 4.{table_number} \u2014 Rekapitulasi Hasil Pengujian Black-Box'
    add_table_caption(doc, recap_caption)

    recap_headers = ['Kode', 'Modul', 'Total TC', 'Berhasil', 'Tidak Berhasil', '% Keberhasilan']
    recap_data = [
        ['F-01', 'Autentikasi', '10', '10', '0', '100%'],
        ['F-02', 'Beranda', '9', '9', '0', '100%'],
        ['F-03', 'Dashboard Analitik', '7', '7', '0', '100%'],
        ['F-04', 'Kasir (Point of Sale)', '26', '26', '0', '100%'],
        ['F-05', 'Manajemen Produk', '17', '17', '0', '100%'],
        ['F-06', 'Laporan', '11', '11', '0', '100%'],
        ['F-07', 'Pengaturan', '29', '29', '0', '100%'],
    ]
    actual_total = sum(int(r[2]) for r in recap_data)  # = 108
    recap_data.append(['', 'TOTAL', str(actual_total), str(actual_total), '0', '100%'])

    tbl_recap = doc.add_table(rows=1 + len(recap_data), cols=6)
    tbl_recap.style = 'Table Grid'
    tbl_recap.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(recap_headers):
        cell = tbl_recap.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_font(run, size=10, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    for r_idx, row in enumerate(recap_data):
        is_total = (r_idx == len(recap_data) - 1)
        for c_idx, val in enumerate(row):
            cell = tbl_recap.rows[r_idx + 1].cells[c_idx]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(val)
            set_font(run, size=10, bold=is_total)
            if is_total:
                set_cell_shading(cell, 'E2EFDA')

    doc.add_paragraph()

    # ============================================================
    # 4. KESIMPULAN
    # ============================================================
    add_heading(doc, 'Kesimpulan', level=3)

    add_para(doc,
        f'Berdasarkan hasil pengujian *black-box* yang telah dilakukan terhadap '
        f'**{actual_total} kasus uji** yang mencakup 7 modul fungsional utama aplikasi '
        f'POS PWA (versi 0.5.0), seluruh kasus uji menghasilkan hasil yang sesuai dengan '
        f'*Hasil Diharapkan* yang telah ditetapkan. Tingkat keberhasilan pengujian mencapai '
        f'**100%**, yang menandakan bahwa sistem telah memenuhi seluruh spesifikasi kebutuhan '
        f'fungsional yang dirancang.')

    add_para(doc,
        'Pengujian mencakup skenario positif (input valid dan alur normal) maupun skenario '
        'negatif (validasi input tidak valid, kondisi *offline*, pembayaran kurang, dan nilai '
        'batas), serta diuji pada dua *viewport* berbeda yaitu Desktop (\u22651024px) dan Mobile '
        '(\u2264768px). Dengan demikian, dapat disimpulkan bahwa sistem POS PWA berfungsi dengan '
        'baik sesuai harapan dan layak untuk digunakan.')

    # Accuracy formula as text
    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_formula.paragraph_format.space_before = Pt(12)
    p_formula.paragraph_format.space_after = Pt(6)
    p_formula.paragraph_format.line_spacing = 1.5

    run_f1 = p_formula.add_run('Akurasi = ')
    set_font(run_f1, italic=True)
    run_f2 = p_formula.add_run(f'(Pengujian Berhasil / Total Pengujian) \u00d7 100%')
    set_font(run_f2)

    p_calc = doc.add_paragraph()
    p_calc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_calc.paragraph_format.line_spacing = 1.5
    run_c = p_calc.add_run(f'= ({actual_total} / {actual_total}) \u00d7 100% = 100%')
    set_font(run_c, bold=True)

    print(f"\n  Total TC parsed: {total_tc}")
    print(f"  Total TC recap:  {actual_total}")


if __name__ == '__main__':
    print('=' * 55)
    print('Generating BAB-BLACKBOX.docx...')
    print(f'Source: {MD_FILE}')
    print('=' * 55)

    doc = setup_document()
    build_document(doc)
    doc.save(OUTPUT_FILE)

    print('=' * 55)
    print(f'[DONE] Saved to: {OUTPUT_FILE}')
    print('=' * 55)
