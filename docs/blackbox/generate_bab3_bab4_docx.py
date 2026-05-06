import os
import re
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    
    # The requirement said "tanpa garis vertikal (hanya garis horizontal pada header dan akhir tabel)"
    # We will remove insideH entirely here, and handle the header bottom line in the cell loop.
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'none')

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'none')

    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'none')

    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')

    tblBorders.append(top)
    tblBorders.append(left)
    tblBorders.append(bottom)
    tblBorders.append(right)
    tblBorders.append(insideH)
    tblBorders.append(insideV)
    
    # Remove existing borders if any
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblPr.append(tblBorders)

def create_omml_equation(p, lhs, numerator, denominator, rhs=''):
    """Create an OMML equation with proper fraction bar in the given paragraph."""
    oMath = OxmlElement('m:oMath')

    def add_math_run(parent, text, style='p'):
        r = OxmlElement('m:r')
        if style:
            rPr = OxmlElement('m:rPr')
            sty = OxmlElement('m:sty')
            sty.set(qn('m:val'), style)
            rPr.append(sty)
            r.append(rPr)
        t = OxmlElement('m:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        parent.append(r)

    # LHS + equals sign
    add_math_run(oMath, lhs + ' = ')

    # Fraction with bar
    frac = OxmlElement('m:f')
    num_el = OxmlElement('m:num')
    add_math_run(num_el, numerator)
    den_el = OxmlElement('m:den')
    add_math_run(den_el, denominator)
    frac.append(num_el)
    frac.append(den_el)
    oMath.append(frac)

    # RHS (e.g., " × 100%")
    if rhs:
        add_math_run(oMath, ' ' + rhs)

    p._p.append(oMath)


def parse_inline(paragraph_or_cell, text):
    if hasattr(paragraph_or_cell, 'add_paragraph'):
        p = paragraph_or_cell.add_paragraph()
    elif hasattr(paragraph_or_cell, 'add_run'):
        p = paragraph_or_cell
    else:
        p = paragraph_or_cell.paragraphs[0]
        
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = p.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    code_parts = re.split(r'(`.*?`)', sub_part)
                    for c_part in code_parts:
                        if c_part.startswith('`') and c_part.endswith('`'):
                            run = p.add_run(c_part[1:-1])
                            run.font.name = 'Courier New'
                        else:
                            br_parts = re.split(r'(<br>)', c_part)
                            for br_part in br_parts:
                                if br_part == '<br>':
                                    p.add_run('\n')
                                elif br_part:
                                    p.add_run(br_part)

def process_markdown_file(input_path, output_path):
    if not os.path.exists(input_path):
        print(f"File not found: {input_path}")
        return

    doc = Document()

    # Set page size to A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # Set margins
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    is_in_table = False
    table_lines = []

    def set_cell_margins(cell, top=0, bottom=0, left=57, right=57):
        """Set cell margins in twentieths of a point (57 twips ≈ 1mm)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)

    def set_cell_vertical_alignment(cell, align='center'):
        """Set vertical alignment of cell content."""
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), align)
        tcPr.append(vAlign)

    def flush_table():
        nonlocal is_in_table, table_lines
        if not table_lines:
            return

        rows = []
        for line in table_lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line[1:-1].split('|')]
                rows.append(cells)

        # Remove separator row (|---|---|)
        if len(rows) > 1 and all(c.replace('-', '').strip() == '' for c in rows[1]):
            rows.pop(1)

        if rows:
            num_cols = max(len(row) for row in rows)
            table = doc.add_table(rows=0, cols=num_cols)

            # Set table width to full content area (14cm = 21cm - 4cm left - 3cm right)
            tbl = table._tbl
            tblPr = tbl.tblPr
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')  # 100% of content area
            for existing in tblPr.findall(qn('w:tblW')):
                tblPr.remove(existing)
            tblPr.append(tblW)

            # Disable autofit so column widths are respected
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            for existing in tblPr.findall(qn('w:tblLayout')):
                tblPr.remove(existing)
            tblPr.append(tblLayout)

            # Determine column width strategy based on content
            # Content area = 14cm = 7938 twips (approx)
            total_width_twips = 7938
            # Heuristic: detect if table has many empty scoring columns (like UAT questionnaire)
            header_row = rows[0] if rows else []
            scoring_keywords = {'SS', 'S', 'CS', 'KS', 'TS'}
            has_scoring = any(h.upper() in scoring_keywords for h in header_row)

            if has_scoring and num_cols >= 5:
                # UAT-style: narrow No, wide Pernyataan, tiny scoring columns
                scoring_count = sum(1 for h in header_row if h.upper() in scoring_keywords)
                non_scoring = num_cols - scoring_count
                scoring_w = 567  # ~1cm per scoring column
                remaining = total_width_twips - (scoring_count * scoring_w)
                if non_scoring == 2:
                    col_widths = [567] + [remaining - 567] + [scoring_w] * scoring_count  # No=1cm, Pernyataan=rest
                else:
                    col_widths = [remaining // non_scoring] * non_scoring + [scoring_w] * scoring_count
            elif num_cols == 2:
                # Two-column tables: equal or proportional
                col_widths = [total_width_twips // 2] * 2
            elif num_cols <= 4:
                col_widths = [total_width_twips // num_cols] * num_cols
            else:
                # Generic: first column narrow, rest equal
                first_w = 567
                rest_w = (total_width_twips - first_w) // (num_cols - 1)
                col_widths = [first_w] + [rest_w] * (num_cols - 1)

            for row_idx, row_data in enumerate(rows):
                row_cells = table.add_row().cells
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(row_cells):
                        cell = row_cells[col_idx]

                        # Set column width
                        if col_idx < len(col_widths):
                            cell.width = col_widths[col_idx]

                        # Parse content
                        parse_inline(cell.paragraphs[0], cell_data)

                        # Format all runs: TNR 10pt, spasi 1
                        for p in cell.paragraphs:
                            p.paragraph_format.line_spacing = 1.0
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                            for r in p.runs:
                                r.font.size = Pt(10)
                                r.font.name = 'Times New Roman'
                                if row_idx == 0:
                                    r.font.bold = True

                        # Header row: centered text
                        if row_idx == 0:
                            for p in cell.paragraphs:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Scoring/narrow columns: center all content
                        if has_scoring and col_idx == 0:
                            for p in cell.paragraphs:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if has_scoring and col_idx >= (num_cols - scoring_count):
                            for p in cell.paragraphs:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Cell margins and vertical alignment
                        set_cell_margins(cell)
                        set_cell_vertical_alignment(cell)

                        # Header bottom border
                        if row_idx == 0:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcBorders = tcPr.first_child_found_in('w:tcBorders')
                            if tcBorders is None:
                                tcBorders = OxmlElement('w:tcBorders')
                                tcPr.append(tcBorders)
                            cell_bottom = OxmlElement('w:bottom')
                            cell_bottom.set(qn('w:val'), 'single')
                            cell_bottom.set(qn('w:sz'), '4')
                            tcBorders.append(cell_bottom)

            set_table_borders(table)
            doc.add_paragraph()

        is_in_table = False
        table_lines = []

    for line in lines:
        line_clean = line.strip()
        
        if line_clean.startswith('|'):
            is_in_table = True
            table_lines.append(line_clean)
            continue
        else:
            if is_in_table:
                flush_table()

        if not line_clean:
            continue
            
        if line_clean.startswith('# '):
            p = doc.add_heading(level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            parse_inline(p, line_clean[2:].upper())
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
        elif line_clean.startswith('## '):
            p = doc.add_heading(level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            parse_inline(p, line_clean[3:])
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
        elif line_clean.startswith('### '):
            p = doc.add_heading(level=3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            parse_inline(p, line_clean[4:])
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
        elif line_clean.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, line_clean[2:])
        elif line_clean.startswith('**Tabel '):
            # Table title: centered, 10pt, bold per pedoman FTI §5.2.6
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            parse_inline(p, line_clean)
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
        elif line_clean.startswith('(Sumber:') or line_clean.startswith('*(Sumber:'):
            # Table source citation: centered, 10pt per pedoman FTI §5.2.6
            clean_src = line_clean.strip('*')
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_src)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        elif re.search(r'\.{5,}\s*\([\d\.x]+\)\s*$', line_clean):
            # Equation block detected by trailing numbering pattern e.g. "............... (3.x)"
            # Format: centered, proper OMML fraction per pedoman FTI §5.2.8
            num_match = re.search(r'\s*(\.{5,}\s*\([\d\.x]+\))\s*$', line_clean)
            eq_number = ''
            eq_body = line_clean
            if num_match:
                eq_number = num_match.group(1).strip()
                eq_body = line_clean[:num_match.start()].strip()

            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

            # Detect fraction pattern: LHS = (Numerator / Denominator) RHS
            frac_match = re.match(r'(.+?)\s*=\s*\((.+?)\s*/\s*(.+?)\)\s*(.*)', eq_body)
            if frac_match:
                lhs = frac_match.group(1).strip()
                numerator = frac_match.group(2).strip()
                denominator = frac_match.group(3).strip()
                rhs = frac_match.group(4).strip()
                create_omml_equation(p, lhs, numerator, denominator, rhs)
            else:
                run = p.add_run(eq_body)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.italic = True

            if eq_number:
                run2 = p.add_run('    ' + eq_number)
                run2.font.size = Pt(12)
                run2.font.name = 'Times New Roman'
                run2.italic = False
        elif line_clean.startswith('**Keterangan:**') or line_clean.startswith('**Keterangan:'):
            # Keterangan block for formula variables: bold label, 10pt, spasi 1
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run('Keterangan:')
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            # If there is content after **Keterangan:**
            remainder = line_clean.replace('**Keterangan:**', '').replace('**Keterangan:', '').strip().rstrip('*')
            if remainder:
                p2 = doc.add_paragraph()
                p2.paragraph_format.line_spacing = 1.0
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
                run2 = p2.add_run(remainder)
                run2.font.size = Pt(10)
                run2.font.name = 'Times New Roman'
        elif line_clean.startswith('![') and '](' in line_clean:
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line_clean)
            if match:
                alt_text = match.group(1)
                img_path_rel = match.group(2)
                md_dir = os.path.dirname(input_path)
                img_path = os.path.normpath(os.path.join(md_dir, img_path_rel))
                
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if os.path.exists(img_path):
                    run = p.add_run()
                    run.add_picture(img_path, width=Cm(14))
                else:
                    run = p.add_run(f"[Gambar tidak ditemukan: {img_path_rel}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                
                if alt_text:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run_cap = p_cap.add_run(alt_text)
                    run_cap.font.italic = True
                    run_cap.font.size = Pt(10)
        else:
            # Check if this is a Keterangan variable line (starts with identifier = description)
            if re.match(r'^\w[\w\s]*=\s', line_clean):
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line_clean)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.27)
                parse_inline(p, line_clean)

    if is_in_table:
        flush_table()

    try:
        doc.save(output_path)
        print(f"Generated {output_path}")
    except PermissionError:
        print(f"ERROR: Permission denied. Please close the file '{output_path}' if it is open in Word, and try again.")

if __name__ == '__main__':
    docs_dir = r'c:\Users\aldyf\Downloads\CODE\antigravity\post-pwa\docs\blackbox'
    docs_dir_root = r'c:\Users\aldyf\Downloads\CODE\antigravity\post-pwa\docs'
    
    bab3_md = os.path.join(docs_dir, 'BAB3-BLACKBOX.md')
    bab3_docx = os.path.join(docs_dir, 'BAB3-BLACKBOX.docx')
    
    bab4_md = os.path.join(docs_dir, 'BAB4-BLACKBOX.md')
    bab4_docx = os.path.join(docs_dir, 'BAB4-BLACKBOX.docx')
    
    bab4_impl_md = os.path.join(docs_dir_root, 'BAB4-IMPLEMENTASI.md')
    bab4_impl_docx = os.path.join(docs_dir_root, 'BAB4-IMPLEMENTASI.docx')
    
    process_markdown_file(bab3_md, bab3_docx)
    process_markdown_file(bab4_md, bab4_docx)
    process_markdown_file(bab4_impl_md, bab4_impl_docx)
