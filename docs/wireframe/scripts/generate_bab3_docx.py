"""
generate_bab3_docx.py
Generator BAB3-Wireframe.docx sesuai Pedoman Penulisan Skripsi FTI 2025-2026.

Referensi pedoman (dari PDF):
  §5.1.3 Marjin  : Kiri 4cm, Atas 3cm, Kanan 3cm, Bawah 3cm
  §5.1.4 Spasi   : Naskah 1.5 spasi (=360 twips), tabel/gambar/pustaka 1 spasi (=240 twips)
  §5.1.5 Paragraf: Alinea baru ketukan ke-6 (~1.25cm), jarak antar paragraf 2 spasi (=480 twips after)
  §5.1.6 Huruf   : TNR 12pt body, TNR 10pt keterangan gambar/tabel (spasi 1)
  §5.2.3 Bab/Sub : Judul BAB = KAPITAL BOLD CENTER; Sub-bab = Bold Rata Kiri Kapital Tiap Kata
  §5.2.7 Gambar  : Keterangan di BAWAH gambar, centered, spasi 1, font 10
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────
BASE   = Path("C:/Users/aldyf/Downloads/CODE/antigravity/post-pwa/docs/wireframe")
MD     = BASE / "WIREFRAME-BAB3.md"
EXPORT = BASE / "export"
OUT    = BASE / "BAB3-Wireframe.docx"

# ── Twips constants (1 cm ≈ 567 twips, 1 inch = 1440 twips) ───
TWIPS_PER_CM  = 567
LINE_1_5      = 360   # 1.5 spasi = 1.5 × 240
LINE_1_0      = 240   # 1 spasi
AFTER_2_SPASI = 480   # 2 spasi jarak antar paragraf (§5.1.5)
INDENT_KE_6   = int(1.25 * TWIPS_PER_CM)   # ketukan ke-6 ≈ 1.25 cm

# ── Width map (cm) per image stem ──────────────────────────────
WIDTH_MAP = {
    'WF-00-D': 14,  'WF-01-D': 14,  'WF-01H-D': 14,  'WF-02-D': 14,
    'WF-03-D': 14,  'WF-04-D': 14,  'WF-05-D': 14,
    'WF-00-M': 5.5, 'WF-01-M': 5.5, 'WF-01H-M': 5.5, 'WF-02-M': 5.5,
    'WF-03-M': 5.5, 'WF-04-M': 5.5, 'WF-05-M': 5.5,
    'WF-00-E': 8,
    'WF-02-MC': 8,  'WF-02-MR': 7.5, 'WF-02-MN': 8,
    'WF-03-MF': 8,  'WF-03-MK': 8,
    'WF-04-MD': 8,
}

def get_width(fname: str) -> float:
    return WIDTH_MAP.get(Path(fname).stem, 11.0)


# ── XML helpers ─────────────────────────────────────────────────

def _set_font(run, size_pt: int = 12):
    """Apply TNR font to a run, handling East Asian fonts too."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    # Force East-Asian compatible font declaration
    rPr = run._r.get_or_add_rPr()
    for tag in ('w:rFonts',):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn('w:ascii'),     'Times New Roman')
        el.set(qn('w:hAnsi'),     'Times New Roman')
        el.set(qn('w:cs'),        'Times New Roman')
        el.set(qn('w:eastAsia'), 'Times New Roman')


def _set_spacing(para, *, before: int = 0, after: int = 0,
                 line: int = LINE_1_5, line_rule: str = 'auto'):
    """
    Set paragraph spacing in twips.
    line_rule='auto'  → proportional (normal use)
    line_rule='exact' → exact line height
    """
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), line_rule)
    old = pPr.find(qn('w:spacing'))
    if old is not None:
        pPr.remove(old)
    pPr.append(sp)


def _set_first_line_indent(para, twips: int = INDENT_KE_6):
    """Set first-line indent (ketukan ke-6)."""
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), str(twips))
    old = pPr.find(qn('w:ind'))
    if old is not None:
        pPr.remove(old)
    pPr.append(ind)


def _clear_indent(para):
    """Remove first-line indent (for headings)."""
    pPr = para._p.get_or_add_pPr()
    old = pPr.find(qn('w:ind'))
    if old is not None:
        pPr.remove(old)


# ── Paragraph builders ──────────────────────────────────────────

def add_bab_heading(doc, text: str):
    """
    Judul BAB: KAPITAL, BOLD, CENTER, TNR 12pt.
    Spacing: 1.5 spasi, no indent, after = 2 spasi (§5.2.3).
    """
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    _set_font(run, 12)
    _set_spacing(p, before=0, after=AFTER_2_SPASI, line=LINE_1_5)
    _clear_indent(p)
    return p


def add_subbab_heading(doc, text: str, level: int = 2):
    """
    Sub-bab & Anak Sub-bab: Bold, Rata Kiri, Kapital Tiap Kata, TNR 12pt.
    Spacing: 1.5 spasi, no indent (§5.2.3).
    """
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    _set_font(run, 12)
    _set_spacing(p, before=0, after=AFTER_2_SPASI, line=LINE_1_5)
    _clear_indent(p)
    return p


def add_body(doc, text: str):
    """
    Paragraf naskah: Justify, TNR 12pt, 1.5 spasi, indent ke-6 (§5.1.4–5.1.6).
    Jarak antar paragraf: 2 spasi after (§5.1.5).
    """
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Inline **bold** / *italic* parsing
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run      = p.add_run(part[2:-2])
            run.bold = True
            _set_font(run, 12)
        elif part.startswith('*') and part.endswith('*'):
            run        = p.add_run(part[1:-1])
            run.italic = True
            _set_font(run, 12)
        elif part:
            run = p.add_run(part)
            _set_font(run, 12)

    _set_spacing(p, before=0, after=AFTER_2_SPASI, line=LINE_1_5)
    _set_first_line_indent(p, INDENT_KE_6)
    return p


def add_caption(doc, text: str):
    """
    Keterangan gambar: CENTER, TNR 10pt, 1 spasi, no indent.
    Ditempatkan di BAWAH gambar (§5.2.7).
    """
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, 10)
    _set_spacing(p, before=0, after=0, line=LINE_1_0)
    _clear_indent(p)
    return p


def add_image_block(doc, img_path: Path, caption_text: str, width_cm: float):
    """
    Gambar centered + keterangan di bawah (§5.2.7).
    Setelah caption ada spacer 2 spasi sebelum paragraf berikutnya.
    """
    # Paragraph gambar
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_indent(p)
    _set_spacing(p, before=0, after=0, line=LINE_1_0)

    run = p.add_run()
    if img_path.exists():
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run = p.add_run(f'[IMAGE NOT FOUND: {img_path.name}]')
        _set_font(run, 10)

    # Keterangan di bawah gambar
    add_caption(doc, caption_text)

    # Spacer paragraph (2 spasi sebelum teks berikutnya)
    sp   = doc.add_paragraph()
    _set_spacing(sp, before=0, after=0, line=AFTER_2_SPASI)
    _clear_indent(sp)


# ── Document setup ──────────────────────────────────────────────

doc = Document()

# Pastikan default style pakai TNR 12
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

# Margins §5.1.3: Kiri 4cm, Atas 3cm, Kanan 3cm, Bawah 3cm
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(4)
sec.right_margin  = Cm(3)
sec.top_margin    = Cm(3)
sec.bottom_margin = Cm(3)


# ── Parse WIREFRAME-BAB3.md ─────────────────────────────────────

with open(MD, encoding='utf-8') as f:
    raw = f.read()

IMG_RE     = re.compile(r'!\[([^\]]+)\]\(export/([^\)]+)\)')
CAP_RE     = re.compile(r'^\*([^*].+[^*])\*\s*$')
HEAD_RE    = re.compile(r'^(#{1,6})\s+(.+)')
HR_RE      = re.compile(r'^-{3,}\s*$')
COMMENT_RE = re.compile(r'^<!--')
BLANK_RE   = re.compile(r'^\s*$')
EMOJI_RE   = re.compile(
    r'[\U0001F000-\U0001FFFF'
    r'\U00002700-\U000027BF'
    r'\U0001F300-\U0001F9FF'
    r'\u2600-\u26FF]+',
    flags=re.UNICODE
)

lines       = raw.split('\n')
i           = 0
pending_img = None  # (Path, alt_text, width_cm)

while i < len(lines):
    line = lines[i]

    # Skip HTML comments and HR
    if COMMENT_RE.match(line) or HR_RE.match(line):
        i += 1
        continue

    # Skip blank lines
    if BLANK_RE.match(line):
        i += 1
        continue

    # ── Headings ──
    m_h = HEAD_RE.match(line)
    if m_h:
        level = len(m_h.group(1))
        text  = m_h.group(2).strip()
        # Strip inline markdown & emojis
        text  = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text  = EMOJI_RE.sub('', text).strip()

        # Flush pending image without caption
        if pending_img:
            add_image_block(doc, pending_img[0], pending_img[1], pending_img[2])
            pending_img = None

        if level == 1:
            add_bab_heading(doc, text)
        else:
            add_subbab_heading(doc, text, level)
        i += 1
        continue

    # ── Image reference ──
    m_img = IMG_RE.search(line)
    if m_img:
        # Flush any previously pending image
        if pending_img:
            add_image_block(doc, pending_img[0], pending_img[1], pending_img[2])
        alt   = m_img.group(1)
        fname = m_img.group(2)
        pending_img = (EXPORT / fname, alt, get_width(fname))
        i += 1
        continue

    # ── Caption (italic line immediately after image) ──
    m_cap = CAP_RE.match(line)
    if m_cap and pending_img:
        cap_text    = m_cap.group(1)
        add_image_block(doc, pending_img[0], cap_text, pending_img[2])
        pending_img = None
        i += 1
        continue

    # ── Body paragraph ──
    stripped = line.strip()
    if stripped and not stripped.startswith('>'):
        # Flush pending image first
        if pending_img:
            add_image_block(doc, pending_img[0], pending_img[1], pending_img[2])
            pending_img = None
        add_body(doc, stripped)

    i += 1

# Flush trailing image
if pending_img:
    add_image_block(doc, pending_img[0], pending_img[1], pending_img[2])


# ── Save & report ───────────────────────────────────────────────

doc.save(str(OUT))

size_kb = OUT.stat().st_size // 1024
d2      = Document(str(OUT))
n_imgs  = len(d2.inline_shapes)
n_paras = len([p for p in d2.paragraphs if p.text.strip()])

print(f"[OK] SAVED  : {OUT}")
print(f"     Size   : {size_kb:,} KB")
print(f"     Paras  : {n_paras}")
print(f"     Images : {n_imgs}")
print()
print("Formatting applied (Pedoman FTI 2025-2026):")
print(f"  Margin      : Kiri 4cm | Atas 3cm | Kanan 3cm | Bawah 3cm")
print(f"  Font body   : TNR 12pt, spasi 1.5 ({LINE_1_5} twips), after {AFTER_2_SPASI} twips")
print(f"  Indent      : First-line {INDENT_KE_6} twips (~1.25 cm, ketukan ke-6)")
print(f"  Font caption: TNR 10pt, spasi 1 ({LINE_1_0} twips)")
print(f"  Caption pos : Di BAWAH gambar, centered")
print(f"  Judul BAB   : KAPITAL BOLD CENTER")
print(f"  Sub-bab     : Bold Rata Kiri")
