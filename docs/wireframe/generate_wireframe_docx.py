"""
Generate BAB3-Wireframe.docx — Otomatis dari data WIREFRAME-BAB3.md
Standar: Pedoman Penulisan Skripsi FTI Unisba Blitar 2025-2026
"""
import os
import re
from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
EXPORT_DIR = os.path.join(SCRIPT_DIR, "export")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "BAB3-Wireframe.docx")

# === WIDTH MAP (cm) by image type ===
def get_width(filename):
    if filename.endswith("-D.png"):
        return 14 if "WF-00" in filename else 15
    if filename.endswith("-M.png"):
        return 6
    if filename in ("WF-02-MC.png", "WF-02-MR.png", "WF-04-MD.png"):
        return 9
    if filename in ("WF-03-MF.png", "WF-03-MK.png"):
        return 8
    return 10

# === DOCUMENT SETUP ===
def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    return doc

# === ADD HEADING ===
def add_heading(doc, text, level=2):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14) if level == 2 else Pt(12)
    run.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return h

# === ADD PARAGRAPH with italic support ===
def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(1.27)
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

# === ADD BOLD PARAGRAPH (for sub-descriptions like "Tampilan Kondisi...") ===
def add_bold_para(doc, bold_text, rest_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.27)
    run_b = p.add_run(bold_text)
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(12)
    parts = re.split(r'(\*[^*]+\*)', rest_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

# === ADD IMAGE + CAPTION ===
def add_image(doc, filename, caption):
    width = get_width(filename)
    path = os.path.join(EXPORT_DIR, filename)
    if not os.path.exists(path):
        print(f"  [SKIP] {filename} not found")
        return
    # Image paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(0)
    run = p_img.add_run()
    run.add_picture(path, width=Cm(width))
    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.line_spacing = 1.0
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(12)
    run_c = p_cap.add_run(caption)
    run_c.font.name = 'Times New Roman'
    run_c.font.size = Pt(10)
    print(f"  [OK] {filename} ({width}cm) -> {caption}")

# ============================================================
# WIREFRAME DATA — extracted from WIREFRAME-BAB3.md
# ============================================================
def build_document(doc):
    # --- 3.5 Main Section ---
    add_heading(doc, "3.5 Perancangan Antarmuka Sistem", level=2)
    add_para(doc, "Perancangan antarmuka merupakan tahapan dalam pengembangan sistem yang bertujuan untuk merencanakan tampilan visual dan pola interaksi yang akan digunakan oleh pengguna akhir. Proses perancangan antarmuka pada penelitian ini mengikuti prinsip *responsive design*, yaitu rancangan yang mampu menyesuaikan tata letak secara otomatis berdasarkan ukuran layar yang digunakan, baik pada komputer maupun pada perangkat ponsel. Pendekatan ini dipilih karena aplikasi dirancang sebagai *Progressive Web App* (PWA) yang dapat diakses melalui berbagai jenis perangkat tanpa memerlukan instalasi terpisah.")
    add_para(doc, "Perancangan antarmuka pada penelitian ini menghasilkan tujuh kelompok halaman utama yang mencakup seluruh alur kerja kasir dan pemilik usaha, yaitu Halaman Login, Halaman Beranda, Halaman Dasbor Analitik, Halaman Kasir, Halaman Manajemen Produk, Halaman Laporan Penjualan, dan Halaman Pengaturan Toko. Setiap halaman dirancang dalam dua tampilan utama, yaitu tampilan komputer dan tampilan ponsel, serta dilengkapi dengan rancangan kondisi atau jendela interaksi yang relevan. Pada tampilan komputer, seluruh halaman utama dapat diakses melalui menu navigasi vertikal (*sidebar*) yang memuat enam item menu, yaitu Beranda, Kasir, Dashboard, Produk, Laporan, dan Pengaturan, sementara pada tampilan ponsel navigasi menggunakan bilah tab horizontal di bagian bawah layar.")

    # --- 3.5.1 Login ---
    add_heading(doc, "3.5.1 Halaman Login", level=3)
    add_para(doc, "Halaman Login adalah halaman pertama yang ditampilkan kepada pengguna sebelum dapat mengakses sistem. Halaman ini menyediakan formulir masuk yang terdiri dari dua kolom isian, yaitu alamat surel dan kata sandi. Setelah pengguna mengisi kedua kolom tersebut dan menekan tombol *Masuk*, sistem akan memeriksa kebenaran data secara otomatis. Apabila data yang dimasukkan tidak sesuai, sistem menampilkan pesan kesalahan di atas formulir sebagai pemberitahuan kepada pengguna.")
    add_para(doc, "Pada tampilan komputer, formulir masuk diletakkan di tengah layar dengan latar belakang berwarna gelap sehingga perhatian pengguna langsung terfokus pada formulir tersebut. Tidak terdapat menu navigasi pada halaman ini karena pengguna belum masuk ke dalam sistem. Setelah proses masuk berhasil, pengguna akan diarahkan secara otomatis ke halaman Beranda.")
    add_image(doc, "WF-00-D.png", "Gambar 3.1 Rancangan Antarmuka Halaman Login \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, tampilan dan susunan formulir tetap sama dengan versi komputer karena desainnya sudah menyesuaikan ukuran layar secara otomatis. Tombol *Masuk* dirancang dengan lebar yang mengisi penuh layar agar mudah ditekan dengan ibu jari.")
    add_image(doc, "WF-00-M.png", "Gambar 3.2 Rancangan Antarmuka Halaman Login \u2014 Tampilan Ponsel")
    add_bold_para(doc, "Tampilan Kondisi Kesalahan Masuk ", "ditampilkan apabila pengguna memasukkan alamat surel atau kata sandi yang tidak sesuai dengan data yang tersimpan dalam sistem. Pada kondisi ini, batas tepi kolom isian berubah menjadi warna merah sebagai tanda visual bahwa data yang dimasukkan tidak valid, dan sebuah pesan kesalahan muncul di atas formulir yang menginformasikan kepada pengguna agar memeriksa kembali data yang dimasukkan. Mekanisme validasi ini berfungsi sebagai lapisan keamanan pertama yang memastikan hanya pengguna yang berwenang yang dapat mengakses sistem.")
    add_image(doc, "WF-00-E.png", "Gambar 3.3 Rancangan Antarmuka Halaman Login \u2014 Kondisi Kesalahan Masuk")

    # --- 3.5.2 Beranda (Home Hub) --- NEW
    add_heading(doc, "3.5.2 Halaman Beranda", level=3)
    add_para(doc, "Halaman Beranda merupakan halaman utama yang ditampilkan setelah pengguna berhasil masuk ke dalam sistem. Halaman ini dirancang sebagai pusat akses cepat (*Quick-Access Hub*) yang mengutamakan kecepatan operasional kasir dengan menempatkan tautan ke halaman-halaman penting di posisi yang mudah dijangkau. Pendekatan ini menerapkan prinsip Hick\u2019s Law untuk meminimalkan waktu pengambilan keputusan melalui pembatasan jumlah pilihan aksi yang terstruktur.")
    add_para(doc, "Pada tampilan komputer, halaman Beranda menampilkan komponen *GreetingBanner* di bagian atas yang memuat sapaan dinamis berdasarkan waktu (Pagi, Siang, Sore, atau Malam), nama toko yang diambil dari basis data, tanggal dan waktu saat ini, serta indikator status koneksi jaringan. Menu navigasi vertikal di sisi kiri layar menampilkan enam item navigasi, yaitu Beranda (aktif), Kasir, Dashboard, Produk, Laporan, dan Pengaturan, sehingga pengguna dapat berpindah halaman kapan saja tanpa kehilangan konteks navigasi. Di bawah banner terdapat dua kartu *MiniStats* yang menampilkan total omzet dan jumlah transaksi hari ini secara langsung dari server. Bagian utama halaman diisi oleh *QuickActionGrid* yang terdiri dari empat kartu aksi, yaitu Mulai Kasir sebagai tombol utama, Kelola Produk, Lihat Laporan, dan Dashboard Analitik, yang masing-masing mengarahkan pengguna ke halaman tujuan yang bersangkutan.")
    add_image(doc, "WF-01H-D.png", "Gambar 3.4 Rancangan Antarmuka Halaman Beranda \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, seluruh komponen Beranda disusun dalam satu kolom vertikal yang dioptimalkan untuk layar sempit. Komponen *GreetingBanner* ditampilkan dalam versi yang lebih ringkas namun tetap mempertahankan hierarki informasi. Kartu *MiniStats* disusun dalam dua kolom berdampingan untuk menghemat ruang. *QuickActionGrid* ditampilkan sebagai daftar vertikal penuh yang memudahkan interaksi menggunakan ibu jari. Menu navigasi utama dipindahkan ke bilah tab di bagian bawah layar dengan tab Beranda aktif di posisi pertama.")
    add_image(doc, "WF-01H-M.png", "Gambar 3.5 Rancangan Antarmuka Halaman Beranda \u2014 Tampilan Ponsel")

    # --- 3.5.3 Dashboard Analitik --- SEPARATED
    add_heading(doc, "3.5.3 Halaman Dasbor Analitik", level=3)
    add_para(doc, "Halaman Dasbor Analitik merupakan halaman yang menyajikan ringkasan kinerja penjualan secara mendalam melalui visualisasi data. Halaman ini dapat diakses melalui kartu aksi *Dashboard* pada Halaman Beranda atau melalui menu navigasi utama. Berbeda dengan Halaman Beranda yang berfokus pada akses cepat, halaman ini mengedepankan prinsip keterlihatan status sistem (*visibility of system status*) dengan menyajikan informasi analitik berupa ringkasan pendapatan harian dan visualisasi tren penjualan secara seketika (*real-time*).")
    add_para(doc, "Pada tampilan komputer, menu navigasi utama ditempatkan di sisi kiri layar dalam bentuk daftar menu vertikal yang memuat enam item, yaitu Beranda, Kasir, Dashboard (aktif), Produk, Laporan, dan Pengaturan. Penempatan item Dashboard pada posisi ketiga setelah Kasir dilakukan untuk menonjolkan peran analitik sebagai bagian inti dari aktivitas operasional, sehingga pemilik usaha dapat langsung memantau kinerja penjualan dari posisi navigasi yang mudah dijangkau. Bagian kanan layar menampilkan kartu ringkasan dan grafik secara berdampingan. Kartu ringkasan disusun dalam empat kolom berjajar yang memuat informasi omzet hari ini, total transaksi, rata-rata nilai per transaksi, dan menu terlaris, sehingga pengguna dapat langsung memantau kinerja usaha tanpa perlu membaca seluruh teks secara terperinci. Di bawah kartu ringkasan terdapat grafik batang yang menampilkan perbandingan pendapatan dalam tujuh hari terakhir serta daftar produk terlaris.")
    add_image(doc, "WF-01-D.png", "Gambar 3.6 Rancangan Antarmuka Halaman Dasbor Analitik \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, susunan halaman berubah menjadi satu kolom vertikal dengan kartu ringkasan ditampilkan terlebih dahulu di bagian atas, diikuti oleh grafik di bawahnya. Menu navigasi utama dipindahkan ke bilah tab yang terletak di bagian paling bawah layar sehingga mudah dijangkau oleh ibu jari pengguna saat menggunakan ponsel dengan satu tangan.")
    add_image(doc, "WF-01-M.png", "Gambar 3.7 Rancangan Antarmuka Halaman Dasbor Analitik \u2014 Tampilan Ponsel")

    # --- 3.5.4 Cashier ---
    add_heading(doc, "3.5.4 Halaman Kasir (Transaksi Penjualan)", level=3)
    add_para(doc, "Halaman Kasir merupakan halaman utama untuk mencatat transaksi penjualan. Halaman ini menampilkan daftar menu dalam bentuk kartu yang memuat foto produk, nama, dan harga. Di bagian atas halaman terdapat pilihan kategori produk yang dapat dipilih untuk menyaring tampilan daftar menu, dan pilihan kategori ini akan tetap terlihat meskipun pengguna menggulir halaman ke bawah.")
    add_para(doc, "Pada tampilan komputer, daftar menu ditampilkan di bagian kiri layar, sementara bagian kanan layar menampilkan daftar pesanan yang sedang berlangsung beserta total harga yang dihitung secara langsung setiap kali ada perubahan. Kasir dapat mengubah jumlah setiap item, menambahkan catatan khusus untuk item tertentu, atau menghapus item dari daftar pesanan langsung dari panel sebelah kanan tanpa perlu berpindah halaman.")
    add_image(doc, "WF-02-D.png", "Gambar 3.8 Rancangan Antarmuka Halaman Kasir \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, daftar menu mengisi seluruh layar agar lebih mudah dibaca. Tombol *Lihat Pesanan* ditampilkan sebagai tombol yang melayang di sudut kanan bawah layar. Apabila tombol tersebut ditekan, panel daftar pesanan akan muncul dari bagian bawah layar sehingga kasir dapat meninjau dan mengubah pesanan tanpa meninggalkan tampilan daftar menu.")
    add_image(doc, "WF-02-M.png", "Gambar 3.9 Rancangan Antarmuka Halaman Kasir \u2014 Tampilan Ponsel")
    add_bold_para(doc, "Tampilan Konfirmasi Pembayaran ", "muncul di atas halaman kasir setelah kasir menekan tombol *Bayar*. Tampilan ini memuat ringkasan pesanan, pilihan metode pembayaran berupa Tunai, QRIS, atau Transfer Bank, serta kolom isian jumlah uang yang diterima dari pelanggan. Untuk mempercepat proses, sistem menyediakan tombol nominal cepat (15K, 20K, 50K, 100K) yang dapat ditekan sebagai pintasan pengisian jumlah uang. Sistem menghitung dan menampilkan jumlah kembalian secara otomatis berdasarkan uang yang dimasukkan oleh kasir.")
    add_image(doc, "WF-02-MC.png", "Gambar 3.10 Rancangan Antarmuka Halaman Kasir \u2014 Tampilan Konfirmasi Pembayaran")
    add_bold_para(doc, "Halaman Nota Pembayaran ", "ditampilkan setelah transaksi berhasil diselesaikan. Nota digital ini menampilkan identitas toko (nama, alamat, nomor telepon, logo), nomor nota, waktu transaksi, nama kasir, rincian item yang dibeli beserta catatan khusus apabila ada, subtotal, total tagihan, metode pembayaran, jumlah uang diterima, serta kembalian. Format tampilan nota mengikuti standar struk kasir pada umumnya. Kasir dapat memilih untuk mencetak nota ke printer yang terhubung melalui Bluetooth atau menutup tampilan apabila pelanggan tidak memerlukan struk fisik.")
    add_image(doc, "WF-02-MR.png", "Gambar 3.11 Rancangan Antarmuka Halaman Kasir \u2014 Tampilan Nota Pembayaran")

    # --- 3.5.5 Product Management ---
    add_heading(doc, "3.5.5 Halaman Manajemen Produk", level=3)
    add_para(doc, "Halaman Manajemen Produk berfungsi sebagai tempat bagi pemilik usaha untuk mengelola data menu yang tersedia di aplikasi. Halaman ini menampilkan daftar seluruh produk dalam bentuk tabel yang memuat informasi nama produk, kategori, harga, dan status ketersediaan. Di bagian atas tabel terdapat kolom pencarian yang memungkinkan pemilik usaha mencari produk tertentu dengan mengetikkan nama produk, dan hasilnya akan muncul secara langsung tanpa perlu menekan tombol tambahan.")
    add_para(doc, "Pada tampilan komputer, tabel daftar produk menampilkan setiap produk dalam satu baris. Di ujung kanan setiap baris terdapat dua tombol aksi, yaitu tombol *Ubah* untuk membuka formulir pengeditan dan tombol *Hapus* untuk menghapus produk dari sistem. Di bagian atas halaman terdapat tombol *Tambah Produk* dan tombol *Kelola Kategori* yang dapat diakses dengan mudah.")
    add_image(doc, "WF-03-D.png", "Gambar 3.12 Rancangan Antarmuka Halaman Manajemen Produk \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, tabel daftar produk ditampilkan dengan baris yang lebih ringkas agar lebih banyak data yang dapat dilihat dalam satu layar. Produk yang ditandai sebagai tidak tersedia ditampilkan dengan warna yang lebih pudar sebagai tanda visual bahwa produk tersebut sedang tidak aktif. Tombol *Ubah* dan *Hapus* tetap tersedia di ujung kanan setiap baris.")
    add_image(doc, "WF-03-M.png", "Gambar 3.13 Rancangan Antarmuka Halaman Manajemen Produk \u2014 Tampilan Ponsel")
    add_bold_para(doc, "Formulir Tambah/Ubah Produk ", "ditampilkan sebagai jendela yang muncul di atas halaman saat tombol *Tambah Produk* atau *Ubah* ditekan. Formulir ini memuat kolom untuk nama produk, kategori, harga jual, area unggah foto produk, serta tombol aktif/nonaktif untuk mengatur status ketersediaan produk. Pemilik usaha tidak perlu mengelola data stok secara angka; cukup mengatur status produk menjadi aktif atau nonaktif untuk menentukan apakah produk tersebut ditampilkan di halaman Kasir.")
    add_image(doc, "WF-03-MF.png", "Gambar 3.14 Rancangan Antarmuka Halaman Manajemen Produk \u2014 Formulir Tambah/Ubah Produk")
    add_bold_para(doc, "Tampilan Kelola Kategori ", "memungkinkan pemilik usaha menambah, mengubah, atau menghapus kelompok kategori produk. Setiap kategori menampilkan jumlah produk yang terhubung sehingga pemilik usaha dapat mengetahui distribusi produk per kategori dengan cepat. Kategori baru dapat ditambahkan langsung melalui kolom isian yang tersedia di bagian bawah jendela.")
    add_image(doc, "WF-03-MK.png", "Gambar 3.15 Rancangan Antarmuka Halaman Manajemen Produk \u2014 Tampilan Kelola Kategori")

    # --- 3.5.6 Sales Report ---
    add_heading(doc, "3.5.6 Halaman Laporan Penjualan", level=3)
    add_para(doc, "Halaman Laporan Penjualan menyajikan riwayat seluruh transaksi yang pernah terjadi dalam sistem. Data transaksi ditampilkan dalam bentuk tabel yang memuat nomor nota, waktu transaksi, metode pembayaran, jumlah item, dan total nominal transaksi. Di bagian atas halaman terdapat tiga kartu indikator yang menampilkan total pendapatan, jumlah transaksi, dan rata-rata nilai transaksi untuk periode yang sedang ditampilkan.")
    add_para(doc, "Pada tampilan komputer, pengguna dapat menyaring data menggunakan pilihan rentang tanggal dan jenis metode pembayaran yang tersedia di bagian atas tabel. Terdapat pula tombol *Ekspor CSV* yang memungkinkan pemilik usaha mengunduh data transaksi dalam format berkas yang dapat dibuka menggunakan aplikasi pengolah angka seperti Microsoft Excel untuk keperluan analisis lebih lanjut. Setiap baris transaksi dilengkapi tombol untuk melihat detail pesanan dan tombol untuk menghapus data transaksi.")
    add_image(doc, "WF-04-D.png", "Gambar 3.16 Rancangan Antarmuka Halaman Laporan Penjualan \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, kartu indikator disusun secara vertikal dengan kartu total pendapatan ditampilkan paling atas. Pilihan penyaringan tanggal dan metode pembayaran ditempatkan dalam dua baris di atas tabel agar tetap dapat diakses dengan mudah. Beberapa kolom tabel yang kurang penting disembunyikan pada tampilan ponsel untuk menghemat ruang layar, namun tetap ditampilkan pada tampilan komputer.")
    add_image(doc, "WF-04-M.png", "Gambar 3.17 Rancangan Antarmuka Halaman Laporan Penjualan \u2014 Tampilan Ponsel")
    add_bold_para(doc, "Tampilan Detail Transaksi ", "muncul saat pengguna menekan tombol lihat detail pada salah satu baris transaksi. Tampilan ini memuat informasi lengkap tentang transaksi tersebut, termasuk nomor nota, waktu, metode pembayaran, daftar item yang dibeli beserta harga satuan, jumlah, dan catatan item apabila ada, serta rincian total pembayaran dan kembalian. Terdapat pula tombol untuk mencetak ulang struk apabila diperlukan.")
    add_image(doc, "WF-04-MD.png", "Gambar 3.18 Rancangan Antarmuka Halaman Laporan Penjualan \u2014 Tampilan Detail Transaksi")

    # --- 3.5.7 Settings (tanpa Backup) ---
    add_heading(doc, "3.5.7 Halaman Pengaturan Toko", level=3)
    add_para(doc, "Halaman Pengaturan Toko merupakan pusat pengelolaan konfigurasi sistem yang terpisah dari proses transaksi. Halaman ini dibagi menjadi beberapa panel konfigurasi yang dapat dipilih melalui menu navigasi, yaitu Informasi Toko, Pengaturan Printer, dan Profil Akun. Pemisahan ini bertujuan agar setiap pengaturan dapat diakses dan dikelola secara terpisah tanpa mengganggu pengaturan lainnya.")
    add_para(doc, "Pada tampilan komputer, menu pilihan panel ditampilkan di sisi kiri layar dalam bentuk daftar menu vertikal, sementara isi dari panel yang dipilih ditampilkan di sisi kanan layar. Panel Informasi Toko memungkinkan pemilik usaha mengisi data identitas toko yang terdiri dari logo, nama toko, alamat lengkap, dan nomor telepon. Data ini digunakan secara otomatis oleh sistem sebagai informasi yang tercetak di bagian atas setiap nota pembayaran. Panel Pengaturan Printer menampilkan status koneksi printer yang terhubung melalui Bluetooth beserta pilihan pengaturan cetak yang dapat diaktifkan atau dinonaktifkan.")
    add_image(doc, "WF-05-D.png", "Gambar 3.19 Rancangan Antarmuka Halaman Pengaturan Toko \u2014 Tampilan Komputer")
    add_para(doc, "Pada tampilan ponsel, pilihan panel ditampilkan dalam bentuk tab berjajar secara mendatar di bagian atas halaman. Pengguna dapat menekan setiap tab untuk berpindah ke panel yang diinginkan. Pengelolaan akun pengguna pada tampilan ponsel diakses melalui panel yang muncul dari bawah layar, yang menampilkan nama pengguna, pilihan untuk memperbarui profil, mengganti kata sandi, dan keluar dari sistem dalam susunan daftar yang mudah dijangkau.")
    add_image(doc, "WF-05-M.png", "Gambar 3.20 Rancangan Antarmuka Halaman Pengaturan Toko \u2014 Tampilan Ponsel")

# === MAIN ===
if __name__ == "__main__":
    print("=" * 50)
    print("Generating BAB3-Wireframe.docx...")
    print(f"Export dir: {EXPORT_DIR}")
    print("=" * 50)

    doc = setup_document()
    build_document(doc)
    doc.save(OUTPUT_FILE)

    print("=" * 50)
    print(f"[DONE] Saved to: {OUTPUT_FILE}")
    print("=" * 50)
